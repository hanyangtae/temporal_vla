#!/usr/bin/env python
"""main.tex 내용을 학회 워드 양식(양식-논문샘플워드.doc)에 채워 main.docx 를 만든다.

양식 구조(LibreOffice 로 docx 변환한 뒤 관찰):
  단락 0~18  : 1단 머리(한글 제목·저자·소속·이메일 / 영문 제목·저자·소속 / 요약)
  단락 19    : 섹션 구분(sectPr) — 이후 2단
  단락 20~35 : 본문(절 제목 bold 9pt, 본문 9pt 들여쓰기, 참고문헌)
템플릿의 pPr/rPr 을 그대로 복제해 글꼴·크기·정렬을 유지하고 텍스트만 바꾼다.

usage:
  python make_docx.py --template <t.docx> --out main.docx
(템플릿 .doc → .docx 변환은 lo-writer 컨테이너: soffice --convert-to docx)
"""
from __future__ import annotations

import argparse
import copy
import re
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt
from docx.oxml.ns import qn

HERE = Path(__file__).resolve().parent
TEX = HERE / "main.tex"
FIGS = HERE / "figs"

# ---------------------------------------------------------------- LaTeX → text
MATH = {
    r"$k=8$": "k = 8", r"$k$": "k", r"$t/T$": "t/T", r"$i$": "i",
    r"$z_i \in \{1,\dots,k\}$": "zᵢ ∈ {1, …, k}", r"$y_i$": "yᵢ",
    r"$I(Z;Y)=\sum_{k}\sum_{c} p(k,c)\log_2\frac{p(k,c)}{p(k)p(c)}$":
        "I(Z;Y) = Σₖ Σ𝒸 p(k,c) log₂ [p(k,c) / (p(k)p(c))]",
    r"$p(k,c)$": "p(k,c)", r"$c$": "c", r"$\sim$": "@SIM@", r"$\pm$": "±", r"$k\ge8$": "k ≥ 8",
    r"$\times$": "×",
}


def detex(s: str) -> str:
    """인라인 LaTeX 를 평문으로. \\textbf 는 **…** 마커로 남겨 run 분리에 쓴다."""
    s = s.replace("\n", " ")
    for k, v in MATH.items():
        s = s.replace(k, v)
    s = re.sub(r"\\textbf\{([^{}]*)\}", r"**\1**", s)
    s = re.sub(r"\\textit\{([^{}]*)\}", r"\1", s)
    s = re.sub(r"Figure~\\ref\{fig:(\w+)\}", lambda m: "Figure " + FIGNUM[m.group(1)], s)
    s = re.sub(r"\\cite\{([^}]*)\}", lambda m: "[" + ",".join(str(CITE[c.strip()]) for c in m.group(1).split(",")) + "]", s)
    s = s.replace("``", "“").replace("''", "”").replace("---", "—").replace("--", "–")
    s = s.replace(r"\&", "&").replace(r"\%", "%").replace("{,}", ",").replace("~", " ")
    s = s.replace(r'H\"aon', "Häon").replace("@SIM@", "~")
    s = re.sub(r"\s+", " ", s).strip()
    assert "\\" not in s and "$" not in s, s
    return s


FIGNUM = {"purity": "1", "readout": "2", "ksweep": "3"}
CITE = {"groot": 1, "robocasa": 2, "saevla": 3, "observing": 4, "egsae": 5, "awe": 6,
        "siglip": 7, "mechinterp": 8, "safe": 9, "lotus": 10, "options": 11}


def parse_tex(src: str):
    """main.tex 에서 요약·절·문단·그림 캡션·참고문헌을 뽑는다."""
    abstract = re.search(r"\\begin\{quote\}\\small\n(.*?)\n\\end\{quote\}", src, re.S).group(1)
    caps = {m.group(1): m.group(2) for m in re.finditer(
        r"\\caption\{(.*?)\}\n\s*\\label\{fig:(\w+)\}", src, re.S)}
    caps = {v: k for k, v in caps.items()}          # label → caption
    body = src[src.index(r"\section{서 론}"):src.index(r"\renewcommand{\refname}")]
    body = re.sub(r"\\begin\{figure\*?\}.*?\\end\{figure\*?\}", "", body, flags=re.S)
    body = re.sub(r"^%.*$", "", body, flags=re.M)
    body = re.sub(r"\\begin\{equation\}.*?\\end\{equation\}", "\n@@EQ@@\n", body, flags=re.S)
    sections = []
    for m in re.finditer(r"\\section\{(.*?)\}\n(.*?)(?=\\section\{|\Z)", body, re.S):
        paras = [p for p in re.split(r"\n\s*\n", m.group(2).strip()) if p.strip()]
        title = m.group(1)
        if all(len(tok) == 1 for tok in title.split()):   # "서 론" → "서론"
            title = title.replace(" ", "")
        sections.append((title, paras))
    refs = re.findall(r"\\bibitem\{\w+\}\s*(.*?)(?=\\bibitem|\\end\{thebibliography\})",
                      src[src.index(r"\begin{thebibliography}"):], re.S)
    return abstract, caps, sections, refs


# ---------------------------------------------------------------- docx helpers
W = qn


def set_text(p, text: str, bold_all=False, size_pt=None):
    """단락의 첫 run 서식을 복제해 텍스트를 다시 채운다. **…** 는 bold run."""
    runs = p.findall(W("w:r"))
    rpr = None
    for r in runs:
        rp = r.find(W("w:rPr"))
        if rp is not None:
            rpr = copy.deepcopy(rp)
            break
    for r in runs:
        p.remove(r)
    for h in p.findall(W("w:hyperlink")):
        p.remove(h)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        bold = part.startswith("**")
        part = part.strip("*") if bold else part
        r = p.makeelement(W("w:r"), {})
        if rpr is not None:
            rp = copy.deepcopy(rpr)
            for tag in ("w:b", "w:bCs"):
                e = rp.find(W(tag))
                if e is not None:
                    rp.remove(e)
            if bold or bold_all:
                rp.insert(0, rp.makeelement(W("w:b"), {}))
            if size_pt is not None:
                for tag in ("w:sz", "w:szCs"):
                    e = rp.find(W(tag))
                    if e is None:
                        e = rp.makeelement(W(tag), {})
                        rp.append(e)
                    e.set(W("w:val"), str(int(size_pt * 2)))
            r.append(rp)
        t = r.makeelement(W("w:t"), {W("xml:space"): "preserve"})
        t.text = part
        r.append(t)
        p.append(r)


def clone_para(proto, text, **kw):
    p = copy.deepcopy(proto)
    set_text(p, text, **kw)
    return p


def picture_para(doc, proto_center, png: Path, width_cm: float):
    """가운데 정렬 단락에 inline 그림 삽입."""
    p = copy.deepcopy(proto_center)
    for r in p.findall(W("w:r")):
        p.remove(r)
    para = doc.add_paragraph()          # python-docx 로 run·그림 생성 후 옮겨 붙임
    run = para.add_run()
    run.add_picture(str(png), width=Cm(width_cm))
    p.append(run._r)
    para._p.getparent().remove(para._p)
    return p


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--template", type=Path, required=True)
    ap.add_argument("--out", type=Path, default=HERE / "main.docx")
    ap.add_argument("--fig1-cm", type=float, default=15.0, help="그림1(전폭) 너비")
    ap.add_argument("--fig-cm", type=float, default=6.8, help="그림2·3(단 내) 너비")
    ap.add_argument("--line", type=int, default=0,
                    help="본문 줄간격 exact(twip/20pt 단위, 예 220=11pt). 0=양식 기본")
    ap.add_argument("--tight", action="store_true",
                    help="머리 빈 문단 제거(제목 위·영문 제목 위·요약 위)")
    ap.add_argument("--keep-blank", action="store_true", help="요약 위 빈 문단 유지")
    ap.add_argument("--ref-pt", type=float, default=8.0, help="참고문헌 글자 크기")
    ap.add_argument("--preview-font", default=None,
                    help="미리보기용: 모든 rFonts 를 이 글꼴로 치환(제출본에는 쓰지 말 것)")
    a = ap.parse_args()

    abstract, caps, sections, refs = parse_tex(TEX.read_text())
    doc = Document(str(a.template))
    body = doc.element.body
    P = [el for el in body if el.tag == W("w:p")]

    # ---- 1단 머리 -------------------------------------------------------
    set_text(P[1], "VLA 모델은 스스로 어떤 action을 수행 중이라고 인지하는가:")
    set_text(P[2], "Auto Encoder 기반 activation clustering을 통한 online action phase 감지")
    set_text(P[3], "박경태, 김상우, 오윤선†")
    set_text(P[4], "한양대학교")
    set_text(P[5], "rudxo1997@hanyang.ac.kr, kimz1121@hanyang.ac.kr, †yoh21@hanyang.ac.kr")
    set_text(P[7], "What Action Does a VLA Model Think It Is Performing?")
    set_text(P[8], "Online Action-Phase Detection via Auto-Encoder-Based Activation Clustering")
    set_text(P[9], "Kyungtae Park, Sangwoo Kim, Yoonseon Oh†")
    set_text(P[10], "Hanyang University")
    set_text(P[14], detex(abstract))
    P[14].find(W("w:pPr")).find(W("w:jc")).set(W("w:val"), "both")
    for el in (P[15], P[16]):
        body.remove(el)
    # 그림 1(전폭)은 1단 영역인 요약 뒤에 둔다 (Word 2단에서 figure* 대응)
    fig1 = picture_para(doc, P[17], FIGS / "fig1_purity_residual_300.png", a.fig1_cm)
    cap1 = clone_para(P[14], "Figure 1: " + detex(caps["purity"]), size_pt=8)
    cap1.find(W("w:pPr")).find(W("w:jc")).set(W("w:val"), "both")
    P[17].addprevious(fig1)
    P[17].addprevious(cap1)
    body.remove(P[18])
    if not a.keep_blank:                          # 영문 소속과 요약 사이 빈 줄 2개
        for el in (P[11], P[13]):
            body.remove(el)
    if a.tight:
        for el in (P[0], P[6], P[11], P[13]):
            body.remove(el)

    # ---- 2단 본문 -------------------------------------------------------
    head_proto, para_proto, ref_proto = P[20], P[21], P[33]
    center_proto = copy.deepcopy(P[29])           # 가운데 정렬 bold 9pt
    sect_break = P[19]
    for el in P[20:36]:
        body.remove(el)
    anchor = body.find(W("w:sectPr"))            # 문서 끝 sectPr 앞에 삽입

    def add(el):
        anchor.addprevious(el)

    def add_figure(label, png, width_cm):
        add(picture_para(doc, center_proto, png, width_cm))
        c = clone_para(para_proto, f"Figure {FIGNUM[label]}: " + detex(caps[label]), size_pt=8)
        ind = c.find(W("w:pPr")).find(W("w:ind"))
        if ind is not None:
            c.find(W("w:pPr")).remove(ind)
        add(c)

    roman = ["Ⅰ", "Ⅱ", "Ⅲ", "Ⅳ"]
    for si, (title, paras) in enumerate(sections):
        add(clone_para(head_proto, f"{roman[si]}. {title}", bold_all=True))
        for para in paras:
            if para.strip() == "@@EQ@@":
                eq = clone_para(center_proto, "margin = I(Z;Y) − I(Z_clock;Y).   (1)")
                for b in eq.iter(W("w:b")):
                    b.getparent().remove(b)
                add(eq)
                continue
            add(clone_para(para_proto, detex(para)))
        if title == "방법":
            add_figure("readout", FIGS / "fig_readout_bar_300.png", a.fig_cm)
        if title == "결과":
            pass
    # 그림 3은 결과 첫 문단 뒤에 넣는 편이 자연스럽지만 Word 흐름상 절 끝에 둔다
    # → 결과 절 안 삽입: 위 루프에서 결과 절 처리 후 위치를 찾아 끼운다
    res_heads = [el for el in body.iter(W("w:p"))
                 if "".join(t.text or "" for t in el.iter(W("w:t"))).startswith("Ⅲ.")]
    first_res_para = res_heads[0].getnext()
    after = first_res_para.getnext().getnext()      # 결과 둘째 문단 뒤
    figp = picture_para(doc, center_proto, FIGS / "fig_k_sweep_300.png", a.fig_cm)
    capp = clone_para(para_proto, "Figure 3: " + detex(caps["ksweep"]), size_pt=8)
    ind = capp.find(W("w:pPr")).find(W("w:ind"))
    if ind is not None:
        capp.find(W("w:pPr")).remove(ind)
    after.addprevious(figp)
    after.addprevious(capp)

    add(clone_para(center_proto, "참 고 문 헌"))
    for i, r in enumerate(refs, 1):
        rp = clone_para(ref_proto, f"[{i}] " + detex(r), size_pt=a.ref_pt)
        sp = rp.find(W("w:pPr")).find(W("w:spacing"))
        if sp is None:
            sp = rp.find(W("w:pPr")).makeelement(W("w:spacing"), {})
            rp.find(W("w:pPr")).append(sp)
        sp.set(W("w:after"), "0"); sp.set(W("w:before"), "0")
        add(rp)

    if a.line:                                    # 본문 줄간격 고정
        for el in body.iter(W("w:p")):
            if el.find(W("w:pPr")) is None or not any(True for _ in el.iter(W("w:drawing"))) and \
               el.find(W("w:pPr")).find(W("w:pStyle")) is not None:
                ppr = el.find(W("w:pPr"))
                if ppr is None:
                    continue
                sp = ppr.find(W("w:spacing"))
                if sp is None:
                    sp = ppr.makeelement(W("w:spacing"), {}); ppr.insert(1, sp)
                sp.set(W("w:line"), str(a.line)); sp.set(W("w:lineRule"), "exact")
    # 2단 섹션을 같은 쪽에서 이어지게(continuous) — 양식 원본은 nextPage
    for s in doc.sections[1:]:
        ty = s._sectPr.find(W("w:type"))
        if ty is None:
            ty = s._sectPr.makeelement(W("w:type"), {})
            s._sectPr.insert(0, ty)
        ty.set(W("w:val"), "continuous")
    if a.preview_font:
        for rf in body.iter(W("w:rFonts")):
            for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
                if rf.get(W(attr)) is not None:
                    rf.set(W(attr), a.preview_font)
    doc.save(str(a.out))
    print("wrote", a.out)


if __name__ == "__main__":
    main()
